"""
生成 Word 报告 — 严格按排版格式要求
用法: python make_report.py
"""
import os, json, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import config

BASE = config.BASE_DIR
CKPT = config.CHECKPOINT_DIR
SAMPLE = config.SAMPLE_DIR

doc = Document()

# ═══════════ 页面设置: A4, 页边距 ═══════════
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ═══════════ 字体辅助函数 ═══════════
def set_font(run, cn_name, en_name='Times New Roman', size=Pt(12), bold=False):
    """设置中英文字体"""
    run.font.name = en_name
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_name)
    rFonts.set(qn('w:ascii'), en_name)
    rFonts.set(qn('w:hAnsi'), en_name)

def add_body(text, indent=True, spacing=1.5):
    """正文段落: 宋体12pt, 1.5倍行距, 首行缩进2字符"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    if indent:
        pf.first_line_indent = Pt(24)  # 2个12pt字符 ≈ 24pt
    run = p.add_run(text)
    set_font(run, '宋体', 'Times New Roman', Pt(12))
    return p

def add_code(text):
    """代码块: Consolas 10.5pt, 单倍行距, 灰色背景"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    # 灰色背景
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
    p._element.get_or_add_pPr().append(shading)
    run = p.add_run(text)
    set_font(run, '宋体', 'Consolas', Pt(10.5))
    return p

def add_poem_lines(lines, caption=''):
    """生成样例: 等宽字体 10.5pt 居中"""
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_font(run, '宋体', 'Consolas', Pt(9))
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for l in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(l)
        set_font(run, '楷体', 'Consolas', Pt(12))

# ═══════════ 标题样式注册 ═══════════
# 报告总标题用单独段落处理, 一级/二级/三级走 Heading 样式
for lv, (size, cn) in [(1, (Pt(16), '黑体')), (2, (Pt(14), '黑体')), (3, (Pt(12), '黑体'))]:
    h = doc.styles[f'Heading {lv}']
    h.font.name = 'Times New Roman'
    h.font.size = size
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.line_spacing = 1.2
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(3)
    h.paragraph_format.first_line_indent = Pt(0)
    rPr = h.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

# ═══════════ 表格函数 ═══════════
_tbl_counter = [0]
_fig_counter = [0]

def add_table(headers, rows, caption=''):
    """表格: 表号上方居中, 宋体10.5pt, 表头加粗, 单倍行距"""
    _tbl_counter[0] += 1
    # 表号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'表{_tbl_counter[0]} {caption}')
    set_font(run, '宋体', 'Times New Roman', Pt(10.5), bold=True)

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(h)
        set_font(run, '宋体', 'Times New Roman', Pt(10.5), bold=True)
        # 表头浅灰背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        c._element.get_or_add_tcPr().append(shading)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]; c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            set_font(run, '宋体', 'Times New Roman', Pt(10.5))
    doc.add_paragraph()
    return t

def add_figure(img_path, caption=''):
    """图片: 图号下方居中"""
    _fig_counter[0] += 1
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f'图{_fig_counter[0]} {caption}')
    set_font(run, '宋体', 'Times New Roman', Pt(10.5))

# ═══════════ 分页 ═══════════
def page_break():
    doc.add_page_break()

# ═══════════ 封面 ═══════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于字符级语言模型的\n七言绝句条件生成')
set_font(run, '黑体', 'Times New Roman', Pt(22), bold=True)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LSTM · GPT-style Transformer · RhymeToneTransformer')
set_font(run, '黑体', 'Times New Roman', Pt(14))

doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据集: DICAlab Ancient Poems → 七言绝句子集 138,500 首\n生成任务: 首句续写 | 藏头诗\n评测指标: PPL | 格式合规率 | 押韵率 | 平仄匹配率 | 语义相关度 | 自重复度')
set_font(run, '宋体', 'Times New Roman', Pt(11))

page_break()

# ═══════════ 一、数据处理 ═══════════
doc.add_heading('一、数据处理', level=1)

doc.add_heading('1.1 数据来源与规模', level=2)
add_body('数据集来源于 DICAlab 发布的 Ancient Poems Dataset，通过 Google Drive 自动下载。'
         '原始数据包含 232,670 首各类古诗，每行为一首诗的纯文本。')

add_table(['属性', '值'], [
    ['原始诗词总数', '232,670'],
    ['七言绝句数', '138,500'],
    ['训练集', '124,650 (90%)'],
    ['验证集', '13,850 (10%)'],
    ['每首 token 数', 'BOS + 7×4 + SEP×3 + EOS = 33'],
    ['序列填充长度', f'MAX_SEQ_LEN = {config.MAX_SEQ_LEN}'],
], '数据集基本统计')

doc.add_heading('1.2 数据预处理', level=2)
add_body('预处理流程分为以下步骤：')
add_body('（1）标点分割：使用中文标点（，。？！；、：）分割每行诗句，去除标点符号，仅保留汉字部分。')
add_body('（2）格式筛选：仅保留恰好 4 句且每句 7 个汉字的严格七言绝句。经筛选后保留 138,500 首。')
add_body('（3）训练/验证划分：随机打乱后按 9:1 比例划分为训练集和验证集（random seed=42）。')

doc.add_heading('1.3 词表构建', level=2)
add_body('采用字符级（character-level）分词方案，每个汉字作为一个独立的 token。'
         '特殊 token 包括 <PAD>（填充）、<BOS>（句首标记）、<EOS>（句尾标记）、'
         '<SEP>（诗句分隔符）和 <UNK>（未知字符）。')

add_table(['Token', '符号', '索引', '用途'], [
    ['PAD', '<PAD>', '0', '批次内变长序列对齐填充'],
    ['BOS', '<BOS>', '1', '标记诗句序列开始'],
    ['EOS', '<EOS>', '2', '标记诗句序列结束'],
    ['SEP', '<SEP>', '3', '分隔四句诗'],
    ['UNK', '<UNK>', '4', '词表外罕见字符'],
], '特殊 Token 定义')

add_body(f'词表构建参数：MAX_VOCAB_SIZE = {config.MAX_VOCAB_SIZE}，MIN_CHAR_FREQ = {config.MIN_CHAR_FREQ}。'
         f'词表按字符频率降序排列，取前 {config.MAX_VOCAB_SIZE} 个且频率 ≥ {config.MIN_CHAR_FREQ} 的字符。'
         f'实际覆盖约 4,996 个汉字。出现频率 < {config.MIN_CHAR_FREQ} 的罕见字符被过滤并映射至 UNK。')

doc.add_heading('1.4 序列化编码', level=2)
add_body('每首七言绝句被编码为固定格式的 token 序列：')
add_code('    <BOS> 句₁ <SEP> 句₂ <SEP> 句₃ <SEP> 句₄ <EOS>')
add_body('总长度约为 33 tokens（1 + 7 + 1 + 7 + 1 + 7 + 1 + 7 + 1）。'
         '训练时统一填充至 MAX_SEQ_LEN=40，超过则截断。'
         '语言建模目标为标准的下一 token 预测：给定 x[:t] 预测 x[t+1]，'
         '仅在非 PAD 位置计算交叉熵损失。')

# ═══════════ 二、模型结构 ═══════════
doc.add_heading('二、模型结构', level=1)

doc.add_heading('2.1 LSTM 字符级语言模型', level=2)
add_body('LSTM 采用 Encoder-free 自回归架构，将诗句生成建模为逐字符的下一 token 预测任务：')

add_table(['组件', '规格', '说明'], [
    ['Embedding', f'd = {config.EMBED_DIM}', 'Xavier Uniform 初始化，gain=1.0'],
    ['LSTM', f'{config.LSTM_LAYERS} 层，hidden = {config.LSTM_HIDDEN}', 'Orthogonal 权重初始化，forget gate bias = 1.0'],
    ['Dropout', f'rate = {config.DROPOUT}', '应用于 Embedding 后、LSTM 层间、Output 前'],
    ['Output', f'Linear({config.LSTM_HIDDEN} → vocab_size)', 'Xavier Uniform 初始化，bias 初始化为 0'],
    ['总参数量', '≈ 8.5M', f'LSTM_HIDDEN 从 384 增至 512 提升容量'],
], 'LSTM 模型结构参数')

doc.add_heading('2.2 GPT-style Transformer', level=2)
add_body('采用 Decoder-only Pre-Norm 架构，包含以下关键技术：')

add_table(['组件', '规格', '说明'], [
    ['Token Embedding', f'd_model = {config.T_D_MODEL}', 'N(0, 0.02) 初始化，与 Output 共享权重'],
    ['位置编码', f'Sinusoidal，max_len = {config.MAX_SEQ_LEN}', '固定正弦位置编码，支持任意 offset 索引'],
    ['Decoder 层数', f'N = {config.T_LAYERS}', '从 6 层降至 5 层以防止过拟合'],
    ['注意力头数', f'nhead = {config.T_NHEAD}', f'head_dim = {config.T_D_MODEL} // {config.T_NHEAD} = {config.T_D_MODEL // config.T_NHEAD}'],
    ['FFN 隐层', f'd_ffn = {config.T_FFN}', 'Linear → GELU → Dropout → Linear，从 1536 降至 1280'],
    ['Weight Tying', 'output.weight = embedding.weight', '输出投影与输入嵌入共享权重，减少参数 + 正则化'],
    ['KV-Cache', '逐层缓存 Key/Value 张量', '推理时增量计算，避免重复编码历史序列'],
    ['总参数量', '≈ 9.8M', '层数 5 + FFN 1280 + 共享权重'],
], 'Transformer 模型结构参数')

doc.add_heading('2.3 韵律监督微调', level=2)
add_body('在预训练 Transformer 上添加两个辅助分类头，通过联合损失进行微调。'
         '辅助头仅在训练时参与梯度和损失计算，推理时仅使用 LM head 进行生成。')

add_table(['辅助头', '结构', '输出', '损失函数', '损失权重'], [
    ['韵母头', 'LN→Linear(h→h)→GELU→Dropout→Linear(h→33)', '33 类韵部', '对比损失', '2.0'],
    ['平仄头', 'LN→Linear(h→h/2)→GELU→Dropout→Linear(h/2→2)', '2 类（平/仄）', '对比损失', '1.0'],
], '辅助分类头结构')

add_body('联合损失函数：Loss = CrossEntropy(LM) + 2.0 × ContrastiveRhyme + 1.0 × ContrastiveTone。'
         '韵母对比损失作用于末字位置（token 6, 14, 22, 30），鼓励正确韵部的概率质量最大化。'
         '平仄对比损失作用于每个非特殊 token 位置，鼓励正确平仄类的概率质量最大化。'
         '微调学习率为 1e-4（远低于预训练的 5e-4），共微调 20 轮。')

doc.add_heading('2.4 初始化策略', level=2)
add_table(['模型组件', '初始化方式', '参数'], [
    ['Embedding (LSTM)', 'Xavier Uniform', 'gain = 1.0'],
    ['LSTM 权重', 'Orthogonal', '正交矩阵初始化'],
    ['LSTM Forget Bias', 'Constant = 1.0', '鼓励长程记忆'],
    ['Embedding (Transformer)', 'Normal', 'μ = 0, σ = 0.02'],
    ['Linear (Transformer)', 'Normal', 'μ = 0, σ = 0.02, bias = 0'],
    ['辅助头 Linear', 'Normal', 'μ = 0, σ = 0.02, bias = 0'],
], '模型初始化策略汇总')

# ═══════════ 三、训练机制 ═══════════
doc.add_heading('三、训练机制与参数', level=1)

doc.add_heading('3.1 训练超参数', level=2)
add_body('模型采用 Adam 优化器配合 StepLR 学习率调度，以下是完整的训练超参数配置：')

add_table(['参数', '值', '说明'], [
    ['BATCH_SIZE', str(config.BATCH_SIZE), '批次大小，动态 padding 对齐'],
    ['LR', str(config.LR), 'Adam 初始学习率'],
    ['WEIGHT_DECAY', str(config.WEIGHT_DECAY), 'Adam 权重衰减（L2 正则化）'],
    ['LR_STEP', str(config.LR_STEP), '学习率衰减周期（每 N epoch）'],
    ['LR_GAMMA', str(config.LR_GAMMA), '学习率衰减因子 γ'],
    ['GRAD_CLIP', str(config.GRAD_CLIP), '梯度裁剪最大范数'],
    ['EPOCHS', str(config.EPOCHS), '最大训练轮数'],
    ['EARLY_STOP_PATIENCE', str(config.EARLY_STOP_PATIENCE), '早停容忍轮数'],
    ['EARLY_STOP_MIN_DELTA', str(config.EARLY_STOP_MIN_DELTA), '改善阈值（PPL 相对下降 0.5%）'],
    ['TRAIN_SPLIT', str(config.TRAIN_SPLIT), '训练/验证集划分比例'],
    ['LSTM_HIDDEN', str(config.LSTM_HIDDEN), 'LSTM 隐藏层维度（384→512）'],
    ['T_LAYERS', str(config.T_LAYERS), 'Transformer 层数（6→5 防过拟合）'],
    ['DROPOUT', str(config.DROPOUT), 'Dropout 比率（0.2→0.25 防过拟合）'],
], '训练超参数汇总表')

doc.add_heading('3.2 早停与断点续训', level=2)
add_body(f'早停机制：当验证集 PPL 在连续 {config.EARLY_STOP_PATIENCE} 轮内相对下降不足'
         f' {config.EARLY_STOP_MIN_DELTA * 100}% 时触发早停。触发时自动恢复至历史最佳 checkpoint。'
         f'训练过程中同时保存 best.pt（最佳 Val PPL 对应模型）和 final.pt（最后一轮模型）。')
add_body('断点续训：训练启动时自动检测已有 checkpoint 文件，若存在则恢复模型权重、'
         '优化器状态、学习率调度器状态和训练历史，从上次中断的 epoch 继续训练。')

doc.add_heading('3.3 学习率调度', level=2)
add_body(f'采用 StepLR 阶梯衰减策略：每 {config.LR_STEP} 轮学习率乘以 {config.LR_GAMMA}。'
         f'实际调度轨迹：lr = 5e-4（epoch 1-{config.LR_STEP - 1}）→ '
         f'2.5e-4（epoch {config.LR_STEP}-{2 * config.LR_STEP - 1}）→ '
         f'1.25e-4（epoch {2 * config.LR_STEP}-{3 * config.LR_STEP - 1}）→ ...')

doc.add_heading('3.4 韵律微调参数', level=2)
add_table(['参数', '值', '说明'], [
    ['FT_EPOCHS', str(config.FT_EPOCHS), '微调训练轮数'],
    ['FT_LR', str(config.FT_LR), '微调学习率（低于预训练一个数量级）'],
    ['FT_RHYME_WEIGHT', str(config.FT_RHYME_WEIGHT), '韵母对比损失在联合损失中的权重'],
    ['FT_TONE_WEIGHT', str(config.FT_TONE_WEIGHT), '平仄对比损失在联合损失中的权重'],
], '韵律微调参数')

# ═══════════ 四、生成机制 ═══════════
doc.add_heading('四、生成机制与采样策略', level=1)

doc.add_heading('4.1 逐句生成与格式保证', level=2)
add_body('生成采用统一的逐句采样策略（_gen_one_line）：每句从模型直接采样恰好 7 个 token，'
         '全程将 SEP 和 EOS 的 logit 设为 -∞ 进行屏蔽，从根本上杜绝了模型自发输出分隔符'
         '导致的行截断问题。首句续写和藏头诗共用同一生成函数，保证行为一致。')
add_body('生成流程：')
add_code('  1. 构建带上下文的 prompt（BOS + 前句 + SEP + 当前提示字）')
add_code('  2. 模型处理 prompt → 逐 token 采样 7 次（SEP/EOS 全程屏蔽）')
add_code('  3. 提取 7 个汉字作为当前行，拼入上下文')
add_code('  4. 重复直至生成全部 4 行')
add_body('此方案避免了复杂的行计数器（cjk_count）和强制换行逻辑（_force_sep），'
         '代码量减少约 40 行，且消除了 LSTM/Transformer 双路径的行为差异。')

doc.add_heading('4.2 采样参数', level=2)
add_body('系统实现了六种采样策略的组合，以下是完整参数配置：')

add_table(['参数', '值', '说明'], [
    ['TEMPERATURE', str(config.TEMPERATURE), '温度缩放：>1 更随机，<1 更确定，=0 退化为贪心'],
    ['TOP_K', str(config.TOP_K), '仅保留概率最高的 K 个候选 token'],
    ['TOP_P', str(config.TOP_P), 'Nucleus Sampling：累积概率达到 p 的最小候选集（0=关闭）'],
    ['REPETITION_PENALTY', str(config.REPETITION_PENALTY), '重复惩罚系数：已出现 token 的 logit /= rp'],
    ['NO_REPEAT_NGRAM_SIZE', str(config.NO_REPEAT_NGRAM_SIZE), '禁止相同 n-gram 再次出现（0=关闭）'],
    ['GEN_NUM', str(config.GEN_NUM), '每条输入生成 N 首候选'],
    ['MAX_GEN_LEN', str(config.MAX_GEN_LEN), '最大生成 token 数量'],
], '采样策略参数汇总')

doc.add_heading('4.3 重复惩罚机制详解', level=2)
add_body(f'重复惩罚（Repetition Penalty, rp={config.REPETITION_PENALTY}）：'
         f'在每次采样前，遍历已生成 token 集合，对每个已出现 token：'
         f'若其 logit > 0 则 logit /= rp；否则 logit *= rp。'
         f'rp > 1 时压制已出现 token，rp = 1 时不生效，rp < 1 时鼓励重复。')
add_body(f'N-gram 禁止（No-Repeat N-gram, n={config.NO_REPEAT_NGRAM_SIZE}）：'
         f'检查当前生成序列末尾 (n-1)-gram 是否与历史中某个 n-gram 的前缀匹配，'
         f'若匹配则将对应下一个 token 的 logit 设为 -∞，从根源上禁止相同 n-gram 再次出现。'
         f'例如 n=3 时，若已生成过 "月影影"，则当序列末尾出现 "月影" 时，'
         f'"影" 被禁止，防止再次形成 "月影影"。'
         f'此两项机制配合使用，有效消除了字符级生成中常见的 "影影影"、"月月" 类坍缩现象。'
         f'同时，重复惩罚仅作用于新生成的 token，prompt（首句/藏头字）中的字不受惩罚，'
         f'避免误伤高频韵脚字导致押韵率下降。')

doc.add_heading('4.4 上下文传递机制', level=2)
add_body('首句续写和藏头诗统一采用带上下文的逐句生成策略。每句生成时，prompt 包含所有已生成前句的完整内容：')
add_code('  句₁: <BOS> 春 → 采样 7 字 → 春XXXXXX')
add_code('  句₂: <BOS> 春XXXXXX <SEP> 花 → 采样 7 字 → 花XXXXXX')
add_code('  句₃: <BOS> 春XXXXXX <SEP> 花XXXXXX <SEP> 秋 → 采样 7 字 → 秋XXXXXX')
add_code('  句₄: <BOS> ... <SEP> 月 → 采样 7 字 → 月XXXXXX')
add_body('此设计使模型在生成每句时都能看到所有前句的末字内容，从而自然协调押韵。'
         '修复前藏头诗因各句独立生成而无上下文，押韵率为 0%；'
         '统一为逐句上下文生成后，首句续写押韵率达 56%~76%，藏头诗达 20%~56%。')

doc.add_heading('4.5 韵部识别机制', level=2)
add_body('押韵检测使用 _get_rhyme_key() 函数去除韵头（i/u/ü）后比较韵腹和韵尾：')
add_code('  ian → an,  uan → an,  üan → an')
add_code('  iang → ang,  uang → ang')
add_code('  ie → e,  ue → e,  üe → e')
add_body('例如 "天"（tian → 韵部 an）与 "山"（shan → 韵部 an）判定为押韵，'
         '比直接用 pypinyin finals 严格字符串匹配更符合传统诗词的押韵习惯。')
add_body('需要说明的是，传统诗词押韵标准为《平水韵》（106 韵部），'
         '其分类体系与基于现代汉语拼音的去韵头规则存在差异。'
         '例如《平水韵》中 "东" 与 "冬" 分属不同韵部，但拼音均为 ong；'
         '"江" 与 "阳" 拼音不同（iang / ang）却可通押。'
         '由于平水韵需要逐字标注韵部，操作复杂度高，'
         '本系统采用了基于 pypinyin 的去韵头简化方案作为押韵判定的近似。'
         '这一简化可能导致部分传统意义上押韵的组合被漏判，'
         '也可能将部分现代同音但传统不同韵的字误判为押韵。')

# ═══════════ 五、训练结果 ═══════════
doc.add_heading('五、训练结果', level=1)

doc.add_heading('5.1 训练数据', level=2)
with open(os.path.join(CKPT, 'lstm_history.json')) as f:
    lstm_hist = json.load(f)
with open(os.path.join(CKPT, 'transformer_history.json')) as f:
    trans_hist = json.load(f)

lstm_best = min(lstm_hist, key=lambda h: h['val_ppl'])
trans_best = min(trans_hist, key=lambda h: h['val_ppl'])

add_table(['模型', '参数量', '最佳 Val PPL', '最佳 Epoch', '总训练 Epochs'], [
    ['LSTM', '≈ 8.5M', f"{lstm_best['val_ppl']:.1f}", str(lstm_best['epoch']), str(len(lstm_hist))],
    ['Transformer', '≈ 9.8M', f"{trans_best['val_ppl']:.1f}", str(trans_best['epoch']), str(len(trans_hist))],
    ['RhymeToneTransformer', '≈ 12.8M', 'Combined Loss: 6.59', '20（微调）', '20'],
], '三模型训练结果对比')

doc.add_heading('5.2 训练曲线', level=2)
for label, fn, cap in [('LSTM', 'lstm_curves.png', 'LSTM 训练曲线（PPL 和 Loss 随 Epoch 变化）'),
                        ('Transformer', 'transformer_curves.png', 'Transformer 训练曲线（PPL 和 Loss 随 Epoch 变化）')]:
    pth = os.path.join(CKPT, fn)
    if os.path.exists(pth):
        add_figure(pth, cap)

doc.add_heading('5.3 训练分析', level=2)
add_body(f'LSTM：Val PPL 从 314.2 持续下降至 {lstm_best["val_ppl"]:.1f}（epoch {lstm_best["epoch"]}），'
         f'50 epoch 内未触发早停，收敛过程平稳但最终 PPL 较高。'
         f'字符级 LSTM 的容量难以完美建模七言诗的复杂分布，但其简单结构作为基线模型已满足基本的格律要求。')
add_body(f'Transformer：Val PPL 从 107.2 快速下降至 {trans_best["val_ppl"]:.1f}（epoch {trans_best["epoch"]}），'
         f'收敛速度远快于 LSTM，最终 PPL 也更低，验证了自注意力机制在捕捉长距离依赖上的优势。'
         f'然而 Train PPL 持续低于 Val PPL 且差距呈扩大趋势，存在轻微过拟合，'
         f'因此将模型层数从 6 降至 5、FFN 隐层从 1536 降至 1280、Dropout 从 0.2 升至 0.25。')
add_body(f'韵律微调：在预训练 Transformer 上额外微调 20 轮，'
         f'Combined Loss 持续下降，韵母 Loss 从 1.50 降至 0.80。'
         f'统一逐句生成后，微调模型在首句续写中押韵率最高（76% vs Transformer 56%），'
         f'但在藏头诗中略低（52% vs 56%），效果不稳定。'
         f'考虑到额外 12.8M 参数和 24 分钟训练成本，建议从默认流程中移除。')

# ═══════════ 六、评测结果 ═══════════
doc.add_heading('六、评测结果', level=1)

doc.add_heading('6.1 指标定义', level=2)
add_table(['评测指标', '定义', '计算方式'], [
    ['困惑度 PPL', 'exp（平均交叉熵）', '验证集逐 token 交叉熵取指数'],
    ['格式合规率', '4 句 × 7 字完全正确的比例', 'len=4 且 all(len(line) = 7)'],
    ['押韵率', '2、4 句末字同韵部且为平声', '去韵头后 rhyme_key 比较 + tone ∈ {1,2}'],
    ['平仄匹配率', '与四种标准格式的最佳匹配', '逐位比较，取 4 种格式中匹配比例最大值'],
    ['语义相关度', 'IDF 加权嵌入余弦相似度 [0, 1+]', '从 checkpoint 提取 Embedding 权重逐字 IDF 加权'],
    ['Bigram 重复率', '2-gram 去重比例', '例："春风吹春风" → bigrams={春风,风吹,吹春,春风} → 1−3/4=0.25'],
    ['Trigram 重复率', '3-gram 去重比例', '例："月影影影" → trigrams={月影影,影影影} → 1−2/2=0'],
    ['行间最大重叠', '任意两行 LCS 长度占比', '例："花落知多少" vs "花落人何在" → LCS="花落"(2)/7=0.29'],
], '评测指标定义与计算方式（附计算示例）')

# 读取评测报告
with open(os.path.join(SAMPLE, 'generation_report.json')) as f:
    report = json.load(f)

def metrics_section(title, data):
    doc.add_heading(title, level=2)
    headers = ['指标'] + [m.upper() for m in data]
    specs = [
        ('样本数', 'count', 'd'),
        ('格式合规率', 'format', '%'),
        ('押韵率', 'rhyme', '%'),
        ('平仄匹配率', 'tone', '%'),
        ('语义相关度', 'relevance', '.3f'),
        ('Bigram 重复 ↓', 'bigram', '.4f'),
        ('Trigram 重复 ↓', 'trigram', '.4f'),
        ('行间最大重叠 ↓', 'overlap', '.4f'),
    ]
    rows = []
    for label, key, fmt in specs:
        row = [label]
        for m in data:
            v = data[m].get(key)
            if v is None: row.append('N/A')
            elif fmt == 'd': row.append(str(int(v)))
            elif fmt == '%': row.append(f'{v:.1%}')
            else: row.append(f'{v:{fmt}}')
        rows.append(row)
    add_table(headers, rows, title)

metrics_section('6.2 首句续写（每模型 25 首）', report['first_line'])
metrics_section('6.3 藏头诗（每模型 25 首）', report['acrostic'])

doc.add_heading('6.4 结果分析', level=2)
add_body('（1）格式合规率：三个模型均达到 100%，证明 SEP/EOS 强制插入与屏蔽机制完全有效，'
         '任何模型在生成时都能严格遵守 4 句 × 7 字的格律要求。')
add_body('（2）押韵率：统一逐句生成后押韵率大幅改善。'
         '首句续写 Transformer 达 56%、RhymeToneTransformer 达 76%；'
         '藏头诗 Transformer 达 56%、RhymeToneTransformer 达 52%。'
         'LSTM 因容量限制押韵率偏低（20%~28%）。'
         '韵律微调模型在首句续写中押韵率最高（76%），但在藏头诗中略低于基础 Transformer（52% vs 56%），'
         '整体而言微调效果不稳定，建议从默认流程中移除。')
add_body('（3）平仄匹配率：三个模型在 60%~83% 之间，高于随机基线 50%，说明模型从训练数据中学习到了部分平仄规律。'
         '但距离严格格律（100%）仍有较大差距，需要显式的格律约束。')
add_body('（4）语义相关度：Transformer（0.73）优于 LSTM（0.56），与 PPL 趋势一致。'
         'IDF 加权机制有效抑制了高频虚字对句子向量的贡献，使语义比较更加聚焦于实义词。')
add_body('（5）自重复指标：Bigram 和 Trigram 重复率均接近 0，'
         '证明了 Repetition Penalty（1.25）+ No-Repeat 3-gram + prompt 不受罚机制有效抑制了字符坍缩问题。')

# ═══════════ 七、生成样例 ═══════════
doc.add_heading('七、生成样例', level=1)

def parse_txt(path, max_g=2, per_g=1):
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    groups = re.split(r'【\d+/\d+】', text)[1:]
    result = []
    for g in groups[:max_g]:
        m_inp = re.search(r'输入:\s*(.+)', g)
        inp = m_inp.group(1).strip() if m_inp else ''
        poems = []
        for m in re.finditer(r'#\d+\s+.*?\n((?:\s{4}.+\n?)+)', g):
            ls = [l.strip() for l in m.group(1).strip().split('\n') if l.strip()]
            if ls: poems.append(ls)
        result.append((inp, poems[:per_g]))
    return result

for idx, (label, fl_f, ac_f) in enumerate([
    ('LSTM', 'lstm_first_line.txt', 'lstm_acrostic.txt'),
    ('Transformer', 'transformer_first_line.txt', 'transformer_acrostic.txt'),
    ('RhymeToneTransformer', 'rhyme_transformer_first_line.txt', 'rhyme_transformer_acrostic.txt'),
]):
    doc.add_heading(f'7.{idx + 1} {label}', level=2)

    doc.add_heading('首句续写样例', level=3)
    fp = os.path.join(SAMPLE, fl_f)
    if os.path.exists(fp):
        for inp, pms in parse_txt(fp, 2, 1):
            for ls in pms:
                add_poem_lines(ls, f'【输入首句】{inp}')
                break

    doc.add_heading('藏头诗样例', level=3)
    fp = os.path.join(SAMPLE, ac_f)
    if os.path.exists(fp):
        for inp, pms in parse_txt(fp, 2, 1):
            for ls in pms:
                tag = '·'.join(inp) if len(inp) == 4 else inp
                add_poem_lines(ls, f'【藏头字】{tag}')
                break

# 综合质量 Top-5
doc.add_heading('7.4 综合质量 Top-5', level=2)
add_body('评分公式：Score = 3.0 × 押韵（0/1）+ 2.0 × 平仄匹配率 + 2.0 × 语义相关度 '
         '+ 1.5 × 字符多样性 − 3 × Bigram 重复 − 5 × Trigram 重复 − 3 × 行间最大重叠。'
         '从全部 150 首生成诗中选出得分最高的 5 首：')

from utils.metrics import poem_score
all_p = []
for mn in ['lstm', 'transformer', 'rhyme_transformer']:
    for mode in ['first_line', 'acrostic']:
        fp = os.path.join(SAMPLE, f'{mn}_{mode}.txt')
        if not os.path.exists(fp): continue
        for inp, pms in parse_txt(fp, 5, 5):
            for ls in pms:
                s = poem_score(ls, None)
                if s > 0: all_p.append((s, mn, inp, ls))
all_p.sort(key=lambda x: -x[0])
for rank, (sc, mn, inp, ls) in enumerate(all_p[:5], 1):
    add_poem_lines(ls, f'【第 {rank} 名】{mn} | 得分 = {sc:.2f} | 输入: {inp}')

# ═══════════ 八、综合分析 ═══════════
doc.add_heading('八、综合分析', level=1)

doc.add_heading('8.1 模型对比', level=2)
add_body('LSTM（PPL 93.5，8.5M 参数）：结构简单、训练收敛稳定，作为字符级基线模型已满足'
         '基本的格律要求（格式合规率 100%）。但语义连贯性明显弱于 Transformer，'
         '部分生成诗句呈现碎片化特征，缺乏整体的诗意逻辑。')
add_body('Transformer（PPL 59.9，9.8M 参数）：PPL 远低于 LSTM，自注意力机制对七言诗的'
         '长距离依赖关系建模出色。首句续写押韵率 56%、藏头诗 56%，语义流畅度最高，'
         '是综合性价比最好的模型。')
add_body('RhymeToneTransformer（Loss 6.59，12.8M 参数）：首句续写押韵率最高（76%），'
         '但藏头诗中略低于 Transformer（52% vs 56%），效果不稳定。'
         '额外 12.8M 参数和 24 分钟训练换来的提升有限，建议移除或改用生成时硬约束替代。')

doc.add_heading('8.2 采样策略影响', level=2)
add_body(f'贪心采样（T = 0）：输出确定性但缺乏创意，容易产生套路化的高频接续句。')
add_body(f'低温采样（T ≤ 0.5，top_k = 10~20）：输出较稳定，多样性有限，适合对质量要求高、'
         f'多样性要求低的场景。')
add_body(f'推荐参数（T = 0.8，top_k = 40，rp = 1.25，ngram = 3）：在创造性与连贯性之间取得最佳平衡，'
         f'是默认推荐配置。')
add_body(f'高温采样（T ≥ 1.2，top_k = 60~80）：随机性增强，可能产生更新颖的词语组合，'
         f'但也增加了不通顺和跑题的风险。')
add_body(f'重复惩罚（rp = 1.25）+ 3-gram 禁止 + prompt 字不受罚：有效消除字符坍缩，同时保护高频韵脚字。')

doc.add_heading('8.3 局限与改进方向', level=2)
add_body('（1）押韵仍需硬约束：统一逐句生成后押韵率提升至 56%~76%，但距离 100% 仍有差距。'
         '最有效的方案是在生成时对末字位置实施词汇约束——'
         '在句₂和句₄末字位置将非同韵、非平声的候选字 logit 设为 -∞，'
         '强制模型在同韵平声字中选择，可将押韵率直接提升至接近 100%。')
add_body('（2）韵律微调效果有限：Loss 层面的改善不转化为生成质量的提升，'
         '建议从默认流程中移除，或改用强化学习方法（如 PPO 以韵律作为奖励信号）。')
add_body('（3）数据规模有限：138K 首七绝对字符级 Transformer 模型而言偏小，'
         '可考虑大规模预训练 + 领域微调的两阶段策略，或引入 BERT 等预训练模型的知识。')
add_body('（4）押韵判定标准简化：由于传统《平水韵》106 韵部需要逐字标注，操作复杂度高，'
         '本系统采用基于 pypinyin 的去韵头方案作为近似。这一简化可能导致部分传统押韵组合被漏判'
         '（如平水韵中可通押但拼音韵母不同的字），实际押韵率可能被低估。未来可引入平水韵字表提升判定精度。')
add_body('（5）评测维度可扩展：当前评测均为自动化指标，可增加人工评估（流畅度、诗意、创意）、'
         'BERTScore、诗歌专用困惑度（PPPL）等指标，以获得更全面的质量评估。')
add_body('（5）藏头诗可改为全自回归：将四个藏头字嵌入 prompt 一次性生成完整四句，'
         '而非当前逐句调用四次模型，可进一步提升押韵率和语义连贯性，同时减少推理耗时。')

# ── 页脚页码 ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 页码域
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)
    set_font(run, '宋体', 'Times New Roman', Pt(9))

# ── 保存 ──
out = os.path.join(BASE, '七言绝句生成系统报告.docx')
doc.save(out)
print(f'报告已保存: {out}')
